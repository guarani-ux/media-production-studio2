import streamlit as st
from openai import OpenAI
from docx import Document
from io import BytesIO

st.set_page_config(page_title="Internal Strategy Engine", page_icon="⚙️", layout="wide")
st.markdown("""<style>#MainMenu {visibility: hidden;} footer {visibility: hidden;} header {visibility: hidden;} .stButton>button {width: 100%; border-radius: 5px; height: 3em; background-color: #1E3A8A; color: white;}</style>""", unsafe_allow_html=True)

client = OpenAI(
    api_key=st.secrets["OPENAI_API_KEY"],
    organization=st.secrets["OPENAI_ORG_ID"],
    project=st.secrets["OPENAI_PROJECT_ID"]
)

if "current_output" not in st.session_state:
    st.session_state.current_output = ""
if "visual_mood" not in st.session_state:
    st.session_state.visual_mood = None
if "history" not in st.session_state:
    st.session_state.history = []

def docx_export(text):
    doc = Document()
    props = doc.core_properties
    props.author = "System Admin"
    doc.add_heading('Strategic Production Brief', 0)
    doc.add_paragraph(text)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

with st.sidebar:
    st.header("🎬 Studio Controls")
    project = st.text_input("Project Name")
    persona = st.radio("Brand Voice", ["Institutional (Board Ready)", "Punchy (Social First)", "Tactical (Field Instructions)"])
    generate_visuals = st.toggle("Generate Visual Storyboard", value=True)
    mode = st.selectbox("Asset Type", ["Strategic Brief", "Technical Shot List", "Narrative Script"])

    if st.button("🚀 EXECUTE PRODUCTION"):
        if project:
            with st.spinner("Synthesizing Strategy..."):
                persona_map = {
                    "Institutional (Board Ready)": "an authoritative CPL executive focused on Policy 3.13.02.",
                    "Punchy (Social First)": "a high-energy viral content creator focused on engagement metrics.",
                    "Tactical (Field Instructions)": "a gritty media producer giving clear, technical on-set directives."
                }
                text_resp = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": f"You are {persona_map[persona]}. Ground everything in the Calgary 2026 Strategy. Format as a clean {mode}."},
                        {"role": "user", "content": project}
                    ]
                )
                st.session_state.current_output = text_resp.choices[0].message.content
                st.session_state.history = [{"role": "assistant", "content": st.session_state.current_output}]

                if generate_visuals:
                    with st.spinner("Generating Mood Board..."):
                        img_resp = client.images.generate(
                            model="dall-e-3",
                            prompt=f"A professional cinematic mood board for a film project about {project}. Style: High-end corporate cinematography, 4k, realistic lighting.",
                            n=1,
                            size="1024x1024"
                        )
                        st.session_state.visual_mood = img_resp.data[0].url
                st.rerun()

if st.session_state.current_output:
    if st.session_state.visual_mood:
        st.image(st.session_state.visual_mood, use_container_width=True)
    st.markdown(st.session_state.current_output)
    st.divider()
    doc_file = docx_export(st.session_state.current_output)
    st.download_button("📥 Download Word Doc", doc_file, "asset_output.docx")
else:
    st.title("STRATEGY ENGINE v2.0")
    st.info("Select project parameters in the sidebar to begin.")
